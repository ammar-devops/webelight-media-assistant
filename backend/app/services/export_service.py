import os
from pathlib import Path

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.platypus import SimpleDocTemplate

from app.core.config import settings


OUTPUT_DIR = Path(settings.OUTPUT_DIR)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_txt(job):

    filename = OUTPUT_DIR / f"{job.job_id}.txt"

    with open(filename, "w", encoding="utf-8") as f:

        f.write("AI MEDIA ASSISTANT\n")
        f.write("=" * 60 + "\n\n")

        f.write(f"Filename : {job.filename}\n")
        f.write(f"Language : {job.language}\n")
        f.write(f"Duration : {job.duration}\n\n")

        f.write("TRANSCRIPT\n")
        f.write("-" * 60 + "\n")
        f.write(job.transcript or "")

        f.write("\n\n")

        f.write("SUMMARY\n")
        f.write("-" * 60 + "\n")
        f.write(job.summary or "")

    return str(filename)


def create_docx(job):

    filename = OUTPUT_DIR / f"{job.job_id}.docx"

    doc = Document()

    doc.add_heading("AI Media Assistant", level=1)

    doc.add_heading("Information", level=2)

    doc.add_paragraph(f"Filename : {job.filename}")
    doc.add_paragraph(f"Language : {job.language}")
    doc.add_paragraph(f"Duration : {job.duration}")

    doc.add_heading("Transcript", level=2)

    doc.add_paragraph(job.transcript or "")

    doc.add_heading("Summary", level=2)

    doc.add_paragraph(job.summary or "")

    doc.save(filename)

    return str(filename)


def create_pdf(job):

    filename = OUTPUT_DIR / f"{job.job_id}.pdf"

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph("<b>AI Media Assistant</b>", styles["Heading1"]))

    story.append(
        Paragraph(
            f"Filename : {job.filename}",
            styles["Normal"],
        )
    )

    story.append(
        Paragraph(
            f"Language : {job.language}",
            styles["Normal"],
        )
    )

    story.append(
        Paragraph(
            f"Duration : {job.duration}",
            styles["Normal"],
        )
    )

    story.append(
        Paragraph("<b>Transcript</b>", styles["Heading2"])
    )

    story.append(
        Paragraph(
            (job.transcript or "").replace("\n", "<br/>"),
            styles["BodyText"],
        )
    )

    story.append(
        Paragraph("<b>Summary</b>", styles["Heading2"])
    )

    story.append(
        Paragraph(
            (job.summary or "").replace("\n", "<br/>"),
            styles["BodyText"],
        )
    )

    pdf = SimpleDocTemplate(str(filename))

    pdf.build(story)

    return str(filename)


def create_srt(job):

    filename = OUTPUT_DIR / f"{job.job_id}.srt"

    with open(filename, "w", encoding="utf-8") as f:

        f.write("1\n")
        f.write("00:00:00,000 --> 00:59:59,000\n")
        f.write(job.transcript or "")

    return str(filename)